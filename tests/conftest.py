import shutil
from pathlib import Path

import pytest
from docx import Document
from pypdf import PdfReader, PdfWriter
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from PIL import Image, ImageDraw, ImageFont

from cost_extractor import ocr_setup

_OLE_MAGIC = bytes.fromhex("D0CF11E0A1B11AE1")


@pytest.fixture
def simple_docx(tmp_path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("Invoice total: $1,234.56 due on receipt.")
    doc.add_paragraph("No money mentioned here.")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Item"
    table.rows[0].cells[1].text = "$500"
    path = tmp_path / "simple.docx"
    doc.save(path)
    return path


@pytest.fixture
def corrupt_docx(tmp_path: Path) -> Path:
    path = tmp_path / "corrupt.docx"
    path.write_bytes(b"not a real docx file, just garbage bytes")
    return path


@pytest.fixture
def password_protected_docx(tmp_path: Path) -> Path:
    # OOXML encryption wraps the document in an OLE Compound File; a real
    # encrypted docx starts with this exact magic number.
    path = tmp_path / "protected.docx"
    path.write_bytes(_OLE_MAGIC + b"\x00" * 64)
    return path


@pytest.fixture
def text_pdf(tmp_path: Path) -> Path:
    path = tmp_path / "text.pdf"
    c = canvas.Canvas(str(path), pagesize=letter)
    c.drawString(72, 700, "Invoice total: $1,234.56 due on receipt.")
    c.drawString(72, 680, "No money mentioned on this line.")
    c.save()
    return path


@pytest.fixture
def scanned_image_pdf(tmp_path: Path) -> Path:
    """A PDF whose page has NO extractable text layer at all — the page
    content is a rendered image of text (as a real scan would produce),
    built via PIL + drawImage with no drawString calls, to genuinely
    exercise the OCR fallback path."""
    img = Image.new("RGB", (1700, 2200), "white")
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("arial.ttf", 48)
    except OSError:
        font = ImageFont.load_default()
    draw.text((100, 100), "Reimbursement due: $2,345.00", fill="black", font=font)
    img_path = tmp_path / "scan_source.png"
    img.save(img_path)

    path = tmp_path / "scanned.pdf"
    c = canvas.Canvas(str(path), pagesize=letter)
    c.drawImage(str(img_path), 0, 0, width=letter[0], height=letter[1])
    c.save()
    return path


@pytest.fixture
def password_protected_pdf(tmp_path: Path) -> Path:
    src = tmp_path / "to_encrypt.pdf"
    c = canvas.Canvas(str(src), pagesize=letter)
    c.drawString(72, 700, "Secret total: $999.00")
    c.save()

    reader = PdfReader(str(src))
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.encrypt(user_password="secret", owner_password="secret")

    path = tmp_path / "protected.pdf"
    with open(path, "wb") as f:
        writer.write(f)
    return path


@pytest.fixture
def tesseract_available() -> bool:
    tess_exe = ocr_setup.get_tesseract_dir() / "tesseract.exe"
    return tess_exe.exists()


@pytest.fixture
def skip_if_no_tesseract(tesseract_available):
    if not tesseract_available:
        pytest.skip("vendor/tesseract/ not populated on this machine")


@pytest.fixture
def scan_image(tmp_path: Path) -> Path:
    """A standalone image of a document, as a phone photo or flatbed scan
    of a receipt would produce — no container, no text layer, OCR only."""
    img = Image.new("RGB", (1700, 600), "white")
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("arial.ttf", 48)
    except OSError:
        font = ImageFont.load_default()
    draw.text((100, 100), "Reimbursement due: $2,345.00", fill="black", font=font)
    path = tmp_path / "receipt.png"
    img.save(path)
    return path


@pytest.fixture
def corrupt_image(tmp_path: Path) -> Path:
    path = tmp_path / "broken.png"
    path.write_bytes(b"not a real png, just garbage bytes")
    return path
