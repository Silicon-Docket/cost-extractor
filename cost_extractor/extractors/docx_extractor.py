"""Extracts text segments (paragraphs + table cells) from a .docx file."""

from __future__ import annotations

import zipfile
from pathlib import Path

import docx
import docx.opc.exceptions

from cost_extractor.extractors.base import ExtractionResult, Status, TextSegment

_OLE_MAGIC = bytes.fromhex("D0CF11E0A1B11AE1")


def _looks_password_protected(path: Path) -> bool:
    if zipfile.is_zipfile(path):
        return False
    with open(path, "rb") as f:
        header = f.read(len(_OLE_MAGIC))
    return header == _OLE_MAGIC


def extract(path: Path) -> ExtractionResult:
    if _looks_password_protected(path):
        return ExtractionResult(
            status=Status.ERROR, error_message="password-protected .docx"
        )

    try:
        document = docx.Document(str(path))
    except (
        docx.opc.exceptions.PackageNotFoundError,
        zipfile.BadZipFile,
        KeyError,
    ) as e:
        return ExtractionResult(
            status=Status.ERROR, error_message=f"corrupt or unreadable .docx: {e}"
        )

    segments: list[TextSegment] = []

    for i, paragraph in enumerate(document.paragraphs, start=1):
        if paragraph.text.strip():
            segments.append(
                TextSegment(text=paragraph.text, location=f"paragraph {i}")
            )

    for t, table in enumerate(document.tables, start=1):
        for r, row in enumerate(table.rows, start=1):
            for c, cell in enumerate(row.cells, start=1):
                if cell.text.strip():
                    segments.append(
                        TextSegment(
                            text=cell.text,
                            location=f"table {t}, row {r}, col {c}",
                        )
                    )

    return ExtractionResult(status=Status.OK, segments=segments)
